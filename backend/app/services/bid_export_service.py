"""
标书导出服务 — 生成 Word 文件
按照政府采购投标文件排版规范：
- 封面：居中大标题 + 供应商信息
- 目录：带序号的章节清单
- 正文：章节标题带编号，正文宋体小四，1.5倍行距，首行缩进
- 附件：PDF 转图片嵌入，图片直接嵌入
- 页边距：上下2.54cm，左右3.17cm（Word默认）
"""
import io
import logging
import os
import re
from datetime import datetime

import fitz  # PyMuPDF — PDF 转图片
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.bid import BidProject, BidSection

logger = logging.getLogger(__name__)

EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "data", "exports")
UPLOAD_BASE = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "data", "uploads")

# 中文数字序号
CN_NUMBERS = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
              '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十',
              '二十一', '二十二', '二十三', '二十四', '二十五']


def _set_cell_font(cell, text, font_name='宋体', font_size=Pt(12), bold=False, alignment=None):
    """设置表格单元格的文字样式"""
    cell.text = ''
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def _add_styled_paragraph(doc, text, font_name='仿宋_GB2312', font_size=Pt(14),
                          bold=False, alignment=None, first_indent=True,
                          line_spacing=1.5, space_after=Pt(0), color=None):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = space_after
    pf.space_before = Pt(0)
    if first_indent:
        pf.first_line_indent = Cm(0.74)  # 两字符缩进
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    return p


def _add_heading(doc, text, level=1):
    """添加自定义章节标题（不用 Word 内置 Heading 样式，避免默认蓝色）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12) if level == 1 else Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    if level == 1:
        # 一级标题：黑体 三号(16pt) 加粗
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.size = Pt(16)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    elif level == 2:
        # 二级标题：黑体 小三(15pt)
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.size = Pt(15)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        # 三级及以下：楷体 四号(14pt) 加粗
        run = p.add_run(text)
        run.font.name = '楷体'
        run.font.size = Pt(14)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _embed_pdf_as_images(doc, pdf_path: str, file_name: str):
    """将 PDF 每页转为图片嵌入 Word 文档"""
    try:
        pdf_doc = fitz.open(pdf_path)
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            img_stream = io.BytesIO(img_bytes)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run()
            run.add_picture(img_stream, width=Inches(6.0))

        pdf_doc.close()
        logger.info(f"PDF嵌入成功: {file_name}, {page_num + 1}页")
    except Exception as e:
        logger.warning(f"PDF嵌入失败: {file_name}, {e}")
        _add_styled_paragraph(doc, f"【附件：{file_name}】（嵌入失败）",
                              color=RGBColor(200, 0, 0), first_indent=False)


def _write_attachment(doc, stripped: str) -> bool:
    """处理附件标记，返回是否命中附件"""
    attachment_match = re.search(r'\[附件:([^:]+):([^\]]+)\]', stripped)
    if not attachment_match:
        return False
    file_path = attachment_match.group(1)
    file_name = attachment_match.group(2)
    before = stripped[:attachment_match.start()].strip()
    if before:
        _add_styled_paragraph(doc, before)
    full_path = os.path.join(UPLOAD_BASE, file_path)
    ext = file_path.rsplit('.', 1)[-1].lower() if '.' in file_path else ''
    if not os.path.exists(full_path):
        _add_styled_paragraph(doc, f"【附件：{file_name}（文件缺失）】",
                              first_indent=False, color=RGBColor(200, 0, 0))
    elif ext in ('jpg', 'jpeg', 'png'):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(full_path, width=Inches(5.5))
        except Exception:
            _add_styled_paragraph(doc, f"【附件：{file_name}（嵌入失败）】", first_indent=False)
    elif ext == 'pdf':
        _embed_pdf_as_images(doc, full_path, file_name)
    else:
        _add_styled_paragraph(doc, f"【附件：{file_name}】", first_indent=False)
    return True


def _add_run_with_inline_format(paragraph, text: str, base_font='仿宋_GB2312', base_size=Pt(14)):
    """解析行内 Markdown（**加粗**、*斜体*、<br> 软换行），添加 run 到段落"""
    # 先按 <br> 拆段，再对每段处理 markdown 加粗/斜体
    segments = re.split(r'<br\s*/?>', text, flags=re.IGNORECASE)
    for seg_idx, segment in enumerate(segments):
        if seg_idx > 0:
            # 段间插入软换行（cell / paragraph 都适用）
            paragraph.add_run().add_break()
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', segment)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = paragraph.add_run(part[1:-1])
                run.font.italic = True
            else:
                run = paragraph.add_run(part)
            run.font.name = base_font
            run.font.size = base_size
            run._element.rPr.rFonts.set(qn('w:eastAsia'), base_font)


def _write_md_table(doc, lines: list):
    """将 Markdown 表格行转为 Word 表格"""
    # 解析表头和数据行（跳过分隔行 |---|）
    rows_data = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')]
        # 去掉首尾空元素
        cells = [c for c in cells if c or cells.index(c) not in (0, len(cells) - 1)]
        if cells and not all(set(c) <= {'-', ':'} for c in cells):
            rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                _add_run_with_inline_format(p, cell_text,
                                            base_font='黑体' if i == 0 else '仿宋_GB2312',
                                            base_size=Pt(11))
                if i == 0:
                    p.runs[0].font.bold = True if p.runs else None
                    # 表头底色
                    shading = cell._element.get_or_add_tcPr()
                    shading_elm = shading.makeelement(qn('w:shd'), {
                        qn('w:fill'): 'F0F0F0', qn('w:val'): 'clear',
                    })
                    shading.append(shading_elm)

    # 表格后空一行
    doc.add_paragraph()


def _write_section_content(doc, content: str):
    """解析章节内容（支持 Markdown 格式），写入 Word 文档"""
    if not content:
        return

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行跳过
        if not stripped:
            i += 1
            continue

        # 附件标记
        if '[附件:' in stripped:
            _write_attachment(doc, stripped)
            i += 1
            continue

        # Markdown 表格（连续 | 开头的行）
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            _write_md_table(doc, table_lines)
            continue

        # Markdown 标题 ## / ### / ####
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            # 去掉标题中的加粗标记
            title = re.sub(r'\*\*([^*]+)\*\*', r'\1', title)
            _add_heading(doc, title, level=min(level, 3))
            i += 1
            continue

        # Markdown 无序列表 - / * / •
        list_match = re.match(r'^[-*•]\s+(.+)$', stripped)
        if list_match:
            text = list_match.group(1)
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_after = Pt(2)
            pf.left_indent = Cm(1.0)
            pf.first_line_indent = Cm(-0.5)  # 悬挂缩进
            run = p.add_run('• ')
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            _add_run_with_inline_format(p, text)
            i += 1
            continue

        # Markdown 有序列表 1. / 2. / 3.
        ordered_match = re.match(r'^(\d+)[.、]\s*(.+)$', stripped)
        if ordered_match:
            num = ordered_match.group(1)
            text = ordered_match.group(2)
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_after = Pt(2)
            pf.left_indent = Cm(1.0)
            pf.first_line_indent = Cm(-0.5)
            run = p.add_run(f'{num}. ')
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            _add_run_with_inline_format(p, text)
            i += 1
            continue

        # 带缩进的子项 （1）/（2）
        sub_match = re.match(r'^[（(]\d+[）)]\s*(.+)$', stripped)
        if sub_match:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_after = Pt(2)
            pf.left_indent = Cm(1.5)
            pf.first_line_indent = Cm(-0.7)
            _add_run_with_inline_format(p, stripped)
            i += 1
            continue

        # 普通段落（含行内 Markdown 加粗/斜体解析）
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_after = Pt(0)
        # 签章行不缩进
        is_sign_line = any(kw in stripped for kw in ['供应商', '盖章', '签章', '日期：', '日　期', '电话：', '地址：', '邮编：', '传真：'])
        if not is_sign_line:
            pf.first_line_indent = Cm(0.74)
        _add_run_with_inline_format(p, stripped)
        i += 1


class BidExportService:

    def _build_doc(self, project_title: str, company_name: str,
                   sections: list, section_map: dict, child_map: dict) -> Document:
        """构建排版规范的投标文件 Word 文档"""
        doc = Document()

        # ── 全局样式设置 ──
        style = doc.styles['Normal']
        style.font.name = '仿宋_GB2312'
        style.font.size = Pt(14)  # 四号
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(0)

        # 页边距设置
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
            # 页眉页脚距离
            section.header_distance = Cm(1.5)
            section.footer_distance = Cm(1.5)

        root_sections = [s for s in sections if s.parent_id is None or s.parent_id not in section_map]

        # ══════════════════════════════════════════════
        #  第一页：封面
        # ══════════════════════════════════════════════
        # 空行撑高
        for _ in range(4):
            doc.add_paragraph()

        # 项目名称
        _add_styled_paragraph(doc, project_title,
                              font_name='方正小标宋简体', font_size=Pt(26),
                              bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              first_indent=False, line_spacing=1.5)

        # 空行
        for _ in range(3):
            doc.add_paragraph()

        # "响应文件" 四个大字
        for char in ['响', '应', '文', '件']:
            _add_styled_paragraph(doc, char,
                                  font_name='方正小标宋简体', font_size=Pt(52),
                                  bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  first_indent=False, line_spacing=1.2,
                                  space_after=Pt(12))

        for _ in range(3):
            doc.add_paragraph()

        # 供应商信息
        from app.config import settings
        _add_styled_paragraph(doc, f"供应商：{company_name}",
                              font_name='楷体', font_size=Pt(16),
                              alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              first_indent=False, space_after=Pt(8))

        today = datetime.now()
        _add_styled_paragraph(doc, f"{today.year} 年 {today.month} 月 {today.day} 日",
                              font_name='楷体', font_size=Pt(16),
                              alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              first_indent=False)

        doc.add_page_break()

        # ══════════════════════════════════════════════
        #  第二页：目录
        # ══════════════════════════════════════════════
        _add_styled_paragraph(doc, "响应文件资料清单",
                              font_name='黑体', font_size=Pt(16),
                              bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              first_indent=False, space_after=Pt(16))

        # 目录表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # 设置表格列宽
        for cell in table.columns[0].cells:
            cell.width = Cm(3)
        for cell in table.columns[1].cells:
            cell.width = Cm(12)

        # 表头
        _set_cell_font(table.rows[0].cells[0], '序号', '黑体', Pt(12), True,
                       WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_font(table.rows[0].cells[1], '资料名称', '黑体', Pt(12), True,
                       WD_ALIGN_PARAGRAPH.CENTER)
        # 表头底色
        for cell in table.rows[0].cells:
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): 'F0F0F0',
                qn('w:val'): 'clear',
            })
            shading.append(shading_elm)

        # 目录行
        for i, s in enumerate(root_sections):
            row = table.add_row()
            cn_num = CN_NUMBERS[i] if i < len(CN_NUMBERS) else str(i + 1)
            _set_cell_font(row.cells[0], cn_num, '仿宋_GB2312', Pt(12), False,
                           WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_font(row.cells[1], s.title, '仿宋_GB2312', Pt(12))

        doc.add_page_break()

        # ══════════════════════════════════════════════
        #  正文章节
        # ══════════════════════════════════════════════
        def write_section(section, index: int, level=1):
            # 每个一级章节前分页（第一个除外由前面的 page_break 处理）
            if level == 1 and index > 0:
                doc.add_page_break()

            # 章节标题（带中文序号）
            if level == 1:
                cn_num = CN_NUMBERS[index] if index < len(CN_NUMBERS) else str(index + 1)
                title_text = f"{cn_num}、{section.title}"
            else:
                title_text = section.title
            _add_heading(doc, title_text, level)

            # 章节内容（支持 Markdown 解析）
            _write_section_content(doc, section.content)

            # 子章节
            for j, child in enumerate(child_map.get(section.id, [])):
                write_section(child, j, level + 1)

        for idx, root in enumerate(root_sections):
            write_section(root, idx)

        # ── 页脚：页码 ──
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 插入页码域代码
            run = p.add_run()
            run.font.size = Pt(10)
            run.font.name = '宋体'
            fld_char_begin = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
            run._element.append(fld_char_begin)
            run2 = p.add_run()
            run2.font.size = Pt(10)
            instr = run2._element.makeelement(qn('w:instrText'), {})
            instr.text = ' PAGE '
            run2._element.append(instr)
            run3 = p.add_run()
            run3.font.size = Pt(10)
            fld_char_end = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
            run3._element.append(fld_char_end)

        return doc

    async def get_project_title(self, db: AsyncSession, project_id: int) -> str:
        """获取项目标题"""
        result = await db.execute(
            select(BidProject).where(BidProject.id == project_id, BidProject.is_deleted == 0)
        )
        project = result.scalar_one_or_none()
        return project.title if project else f"标书_{project_id}"

    async def export_to_word_stream(self, db: AsyncSession, project_id: int) -> io.BytesIO:
        """导出标书为 Word，返回内存流"""
        project_result = await db.execute(
            select(BidProject).where(BidProject.id == project_id, BidProject.is_deleted == 0)
        )
        project = project_result.scalar_one_or_none()
        if not project:
            raise Exception("标书项目不存在")

        sections_result = await db.execute(
            select(BidSection).where(
                BidSection.project_id == project_id, BidSection.is_deleted == 0
            ).order_by(BidSection.sort_order.asc(), BidSection.id.asc())
        )
        sections = sections_result.scalars().all()

        section_map = {s.id: s for s in sections}
        child_map = {}
        for s in sections:
            if s.parent_id and s.parent_id in section_map:
                child_map.setdefault(s.parent_id, []).append(s)

        from app.config import settings
        doc = self._build_doc(project.title, settings.COMPANY_NAME, sections, section_map, child_map)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    async def export_to_word(self, db: AsyncSession, project_id: int) -> str:
        """导出标书为 Word 文件，返回文件路径"""
        buf = await self.export_to_word_stream(db, project_id)
        title = await self.get_project_title(db, project_id)

        os.makedirs(EXPORT_DIR, exist_ok=True)
        filename = f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        safe_filename = "".join(c for c in filename if c.isalnum() or c in ('_', '-', '.', '（', '）')).strip()
        if not safe_filename.endswith('.docx'):
            safe_filename += '.docx'
        filepath = os.path.join(EXPORT_DIR, safe_filename)

        with open(filepath, 'wb') as f:
            f.write(buf.getvalue())
        return filepath


bid_export_service = BidExportService()
