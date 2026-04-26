"""
文档解析服务 — PDF/Word 文本提取
"""
import logging
import os
import time
from pathlib import Path

import fitz  # PyMuPDF
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class DocumentParser:

    def extract_text(self, file_path: str) -> dict:
        """根据文件类型提取文本，返回 {text, page_count, tables_md, timing}"""
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return self.extract_from_pdf(file_path)
        elif ext == ".docx":
            return self.extract_from_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    def extract_from_pdf(self, file_path: str) -> dict:
        """PDF 文本提取：PyMuPDF 提文字 + pdfplumber 提表格，按页内联"""
        # 先用 pdfplumber 提每页的 markdown 表格
        t0 = time.perf_counter()
        tables_by_page: dict[int, list[str]] = {}
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_tables = page.extract_tables()
                    if not page_tables:
                        continue
                    md_list = [md for t in page_tables if (md := self._table_to_markdown(t))]
                    if md_list:
                        tables_by_page[i] = md_list
        except Exception as e:
            logger.warning(f"[TIMING] pdfplumber 表格提取失败: {e}")
        t_tables = time.perf_counter() - t0

        # PyMuPDF 提取文字，并把同页 markdown 表格紧跟其后
        t1 = time.perf_counter()
        doc = fitz.open(file_path)
        pages_text = []
        all_tables_md = []
        for i, page in enumerate(doc):
            text = page.get_text("text")
            chunk = f"--- 第 {i+1} 页 ---\n{text}"
            for md in tables_by_page.get(i, []):
                chunk += f"\n\n【本页表格】\n{md}"
                all_tables_md.append(f"[第{i+1}页表格]\n{md}")
            pages_text.append(chunk)
        page_count = len(doc)
        doc.close()
        t_text = time.perf_counter() - t1
        t_total = time.perf_counter() - t0

        full_text = "\n\n".join(pages_text)
        logger.info(
            f"[TIMING] PDF 提取完成: 总 {t_total:.2f}s | pdfplumber表格 {t_tables:.2f}s | "
            f"PyMuPDF文字 {t_text:.2f}s | 页数 {page_count} | 字数 {len(full_text)} | 表格 {len(all_tables_md)}"
        )
        return {
            "text": full_text,
            "page_count": page_count,
            "tables_md": all_tables_md,
            "timing": {
                "extract_total": round(t_total, 2),
                "extract_tables": round(t_tables, 2),
                "extract_text": round(t_text, 2),
            },
        }

    def extract_from_docx(self, file_path: str) -> dict:
        """Word 文本提取"""
        t0 = time.perf_counter()
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                # 标题层级
                if "Heading" in (para.style.name or ""):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        level_num = int(level)
                        paragraphs.append(f"{'#' * level_num} {para.text}")
                    except ValueError:
                        paragraphs.append(para.text)
                else:
                    paragraphs.append(para.text)

        # 提取表格
        tables_md = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            md = self._table_to_markdown(rows)
            if md:
                tables_md.append(f"[表格{i+1}]\n{md}")

        full_text = "\n".join(paragraphs)
        if tables_md:
            full_text += "\n\n=== 表格数据 ===\n\n" + "\n\n".join(tables_md)

        # Word 没有准确页数，估算
        page_count = max(1, len(full_text) // 2000)
        t_total = time.perf_counter() - t0
        logger.info(
            f"[TIMING] DOCX 提取完成: {t_total:.2f}s | 估算页数 {page_count} | 字数 {len(full_text)} | 表格 {len(tables_md)}"
        )
        return {
            "text": full_text,
            "page_count": page_count,
            "tables_md": tables_md,
            "timing": {"extract_total": round(t_total, 2)},
        }

    def _table_to_markdown(self, table: list) -> str:
        """表格转 Markdown（cell 内换行转 <br>，否则一个 cell 会跨行打散表格）"""
        if not table or not table[0]:
            return ""

        def _clean_cell(cell) -> str:
            s = str(cell or "").strip()
            # cell 内换行转 <br>，避免 markdown 把 cell 内容当成新行解析
            return s.replace("\r\n", "\n").replace("\n", "<br>")

        cleaned = [[_clean_cell(c) for c in row] for row in table]
        cleaned = [row for row in cleaned if any(cell for cell in row)]
        if not cleaned:
            return ""

        header = "| " + " | ".join(cleaned[0]) + " |"
        separator = "| " + " | ".join("---" for _ in cleaned[0]) + " |"
        rows = ["| " + " | ".join(row) + " |" for row in cleaned[1:]]
        return "\n".join([header, separator] + rows)


document_parser = DocumentParser()
